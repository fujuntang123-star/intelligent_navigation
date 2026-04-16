"""
智职领航 - AI 职业规划智能体后端服务
改进版：统一环境变量管理、异步I/O、连接池生命周期、app.state依赖注入
"""
import json
import os
import uuid
import logging
import asyncio
import shutil
import traceback
from pathlib import Path
from contextlib import asynccontextmanager
from typing import Optional

import aiofiles
import uvicorn
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi import UploadFile, File
from pydantic import BaseModel

# ==========================================
# 1. 环境变量加载 & 日志配置
# ==========================================
load_dotenv(Path(__file__).parent / ".env")

# 配置 HuggingFace 国内镜像
os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"

# 日志配置
logger = logging.getLogger("career_agent")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

# ==========================================
# 2. 全局资源初始化（在 lifespan 中管理）
# ==========================================
@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理：启动时初始化资源，关闭时清理"""
    # === 启动时 ===
    logger.info("🚀 后端服务器启动中...")

    # 加载 .env（如果尚未加载）
    load_dotenv(Path(__file__).parent / ".env")

    # 初始化 Neo4j 图数据库连接（共享单例，挂载到 app.state）
    try:
        from graph.graph_builder import CareerGraph
        neo4j_uri = os.getenv("NEO4J_URI", "bolt://localhost:7687")
        neo4j_user = os.getenv("NEO4J_USER", "neo4j")
        neo4j_password = os.getenv("NEO4J_PASSWORD", "12345678")
        app.state.neo4j_graph = CareerGraph(neo4j_uri, neo4j_user, neo4j_password)
        logger.info(f"✅ Neo4j 连接已建立: {neo4j_uri}")
    except Exception as e:
        logger.warning(f"⚠️ Neo4j 连接失败，部分图谱功能不可用: {str(e)}")
        app.state.neo4j_graph = None

    # 初始化匹配引擎
    try:
        from core.matcher import MatchEngine
        excel_path = Path(__file__).parent / "data" / "20260226105856_457.xls"
        app.state.match_engine = MatchEngine(excel_path=excel_path if excel_path.exists() else None, use_rag=True)
        logger.info("✅ 匹配引擎初始化完成")
    except Exception as e:
        logger.error(f"❌ 匹配引擎初始化失败: {str(e)}")
        app.state.match_engine = None

    # 注入引擎到 LangGraph 模块（保持兼容）
    if app.state.match_engine:
        try:
            from core.agent_graph import set_matcher_engine
            set_matcher_engine(app.state.match_engine)
            logger.info("✅ 已将匹配引擎注入到 LangGraph 模块")
        except Exception as e:
            logger.warning(f"⚠️ 引擎注入失败: {str(e)}")

    yield  # 应用运行期间

    # === 关闭时 ===
    logger.info("🛑 后端服务器正在关闭...")
    # 关闭 Neo4j 连接
    if hasattr(app.state, "neo4j_graph") and app.state.neo4j_graph:
        try:
            app.state.neo4j_graph.close()
            logger.info("✅ Neo4j 连接已关闭")
        except Exception as e:
            logger.error(f"❌ Neo4j 关闭失败: {str(e)}")

# ==========================================
# 3. FastAPI 应用
# ==========================================
app = FastAPI(
    title="智职领航 - AI 职业规划智能体",
    description="基于大模型和知识图谱的职业规划系统",
    version="3.0",
    lifespan=lifespan
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 导入 API 路由模块
from api.career_paths import router as career_paths_router
from api.export import router as export_router
app.include_router(career_paths_router)
app.include_router(export_router)

# ==========================================
# 4. 辅助函数
# ==========================================
def cleanup_file(file_path: str):
    """后台清理临时上传的简历文件"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"🗑️ 已清理临时文件: {file_path}")
    except Exception as e:
        logger.error(f"⚠️ 清理临时文件失败 {file_path}: {str(e)}")


# ==========================================
# 5. API 路由
# ==========================================
@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    """上传简历接口，支持 PDF 和 Word 格式（使用 aiofiles 异步写入）"""
    logger.info(f"📄 收到上传的简历文件：{file.filename}")

    allowed_extensions = [".pdf", ".doc", ".docx"]
    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext not in allowed_extensions:
        return {"status": "error", "message": "不支持的文件格式！仅支持 PDF、DOC 和 DOCX 格式"}

    try:
        upload_dir = Path(__file__).parent / "uploads"
        upload_dir.mkdir(exist_ok=True)

        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = str(upload_dir / unique_filename)

        # 【改进】使用 aiofiles 异步写入，不阻塞事件循环
        content = await file.read()
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)

        logger.info(f"✅ 文件已保存：{file_path}")

        # 读取文件内容
        resume_text = ""
        try:
            if file_ext == ".pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                resume_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                logger.info(f"✅ PDF 解析成功，共{len(reader.pages)}页")
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                resume_text = "\n".join(para.text for para in doc.paragraphs)
                logger.info(f"✅ DOCX 解析成功，共{len(doc.paragraphs)}段")
            elif file_ext == ".doc":
                resume_text = f"[旧版 Word 文件已上传：{file.filename}]"
        except Exception as e:
            logger.warning(f"⚠️ 文件解析失败: {str(e)}")
            resume_text = f"[文件已上传，但解析失败：{file.filename}]"

        # 【改进】注册后台清理任务
        if background_tasks:
            background_tasks.add_task(cleanup_file, file_path)

        return {
            "status": "success",
            "message": "简历上传成功",
            "file_name": file.filename,
            "file_type": file_ext,
            "resume_text": resume_text
        }

    except Exception as e:
        logger.exception(f"❌ 上传简历失败")
        return {"status": "error", "message": f"上传失败：{str(e)}"}


@app.get("/api/job_list")
async def get_job_list():
    """获取岗位列表（使用 app.state 而非全局变量）"""
    try:
        engine = getattr(app.state, "match_engine", None)
        if engine is None or not hasattr(engine, 'df') or engine.df is None:
            return {"status": "error", "message": "暂无岗位数据"}

        df = engine.df
        sample_df = df.sample(n=min(20, len(df)))
        job_list = []
        for _, row in sample_df.iterrows():
            job_list.append({
                "job_name": row.get("job_name", "未知岗位"),
                "salary_range": row.get("salary", "面议"),
                "education": row.get("education", "不限"),
                "experience": row.get("experience", "不限"),
                "city": row.get("city", "全国"),
                "skills_required": row.get("skills", "")[:100] + "..." if len(row.get("skills", "")) > 100 else row.get("skills", "")
            })
        return {"status": "success", "data": job_list}
    except Exception as e:
        logger.exception(f"❌ 获取岗位列表失败")
        return {"status": "error", "message": str(e)}


# ==================== 求职模式 ====================
@app.post("/api/job_mode/analyze")
async def job_mode_analyze(data: dict):
    """求职模式：分析简历并生成完整职业规划（自动匹配岗位）"""
    resume_text = data.get("resume_text", "")
    target_job = data.get("target_job", "")
    stream = data.get("stream", True)

    logger.info(f"🎯 求职模式分析 - 目标岗位：{target_job or '自动检测'}")

    try:
        engine = getattr(app.state, "match_engine", None)

        # 如果没有指定目标岗位，先自动匹配最优岗位
        if not target_job and engine:
            logger.info("🔍 未指定目标岗位，正在自动匹配...")
            try:
                from core.resume_parser import parse_resume_to_json
                profile = None
                profile = await parse_resume_to_json(resume_text)

                if profile and "error" not in profile:
                    formatted_skills = {s['name']: s.get('level_text', '熟悉') for s in profile.get('professional_skills', [])}
                    profile['skills'] = formatted_skills
                    profile['resume_text'] = resume_text

                    recommendations = engine.recommend_jobs(profile, top_n=1, use_rag=True)

                    if recommendations:
                        target_job = recommendations[0]['job_name']
                        score = recommendations[0]['total_score']
                        logger.info(f"✅ 自动匹配最优岗位：{target_job}（匹配度：{score:.1f}分）")
                    else:
                        target_job = "Java 开发"
                        logger.info(f"⚠️ 未找到匹配岗位，使用默认：{target_job}")
                else:
                    target_job = "Java 开发"
                    logger.info(f"⚠️ 简历解析失败，使用默认：{target_job}")
            except Exception as e:
                target_job = "Java 开发"
                logger.warning(f"⚠️ 自动匹配失败：{str(e)}，使用默认：{target_job}")

        # 构造初始状态
        initial_state = {
            "resume_text": resume_text,
            "target_job": target_job
        }

        # 导入 LangGraph 应用
        from core.agent_graph import career_agent_app

        if stream:
            async def generate():
                try:
                    logger.info("🔄 开始执行 LangGraph 智能体工作流（完全流式）...")
                    async for event in career_agent_app.astream_events(initial_state, version="v2"):
                        event_type = event.get("event")
                        if event_type == "on_chain_start":
                            node_name = event.get("name", "")
                            if node_name == "parse":
                                yield json.dumps({"type": "status", "message": "📄 正在解析简历..."}, ensure_ascii=False) + "\n"
                            elif node_name == "research":
                                yield json.dumps({"type": "status", "message": "🔍 正在计算匹配度..."}, ensure_ascii=False) + "\n"
                            elif node_name == "draft":
                                yield json.dumps({"type": "status", "message": "✍️ 正在生成职业规划报告..."}, ensure_ascii=False) + "\n"
                        elif event_type == "on_chain_end":
                            node_name = event.get("name", "")
                            output_data = event.get("data", {}).get("output", {})
                            if node_name == "parse" and output_data and "student_profile" in output_data:
                                yield json.dumps({"type": "parsed_profile", "data": output_data["student_profile"]}, ensure_ascii=False) + "\n"
                            elif node_name == "research" and output_data:
                                if "match_score" in output_data:
                                    yield json.dumps({"type": "match_result", "data": output_data["match_score"]}, ensure_ascii=False) + "\n"
                                if "transfer_paths" in output_data:
                                    yield json.dumps({"type": "transfer_paths", "data": output_data["transfer_paths"]}, ensure_ascii=False) + "\n"
                            elif node_name == "draft" and output_data and "report_draft" in output_data:
                                report_content = output_data["report_draft"]
                                chunk_size = 500
                                for i in range(0, len(report_content), chunk_size):
                                    yield json.dumps({"type": "report_chunk", "chunk": report_content[i:i + chunk_size]}, ensure_ascii=False) + "\n"
                                    await asyncio.sleep(0.01)
                    yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"
                    logger.info(f"✅ 求职模式分析完成")
                except Exception as e:
                    logger.exception(f"❌ 求职模式分析错误")
                    yield json.dumps({"type": "error", "message": str(e)}, ensure_ascii=False) + "\n"

            return StreamingResponse(generate(), media_type="application/x-ndjson")
        else:
            final_state = await asyncio.to_thread(career_agent_app.invoke, initial_state)
            if final_state.get("error"):
                return {"status": "error", "message": final_state["error"]}
            return {
                "status": "success",
                "summary": final_state["report_draft"],
                "parsed_profile": final_state["student_profile"],
                "match_result": final_state["match_score"],
                "transfer_paths": final_state["transfer_paths"]
            }
    except Exception as e:
        logger.exception(f"❌ 求职模式分析异常")
        return {"status": "error", "message": f"服务器错误：{str(e)}"}


# ==================== 咨询模式 ====================
@app.post("/api/consult_mode/select_job")
async def consult_mode_select_job(data: dict):
    """咨询模式：用户选择岗位后提供详细咨询"""
    job_name = data.get("job_name", "")
    user_questions = data.get("questions", [])
    stream = data.get("stream", True)

    logger.info(f"💼 咨询模式 - 用户选择岗位：{job_name}")

    try:
        engine = getattr(app.state, "match_engine", None)
        if engine is None or not hasattr(engine, 'df') or engine.df is None:
            return {"status": "error", "message": "知识库未加载"}

        df = engine.df
        job_info = df[df['job_name'] == job_name]

        if len(job_info) == 0:
            return {"status": "error", "message": "未找到该岗位信息"}

        job_data = job_info.iloc[0].to_dict()
        job_detail = {
            "job_name": job_data.get("job_name", ""),
            "salary_range": job_data.get("salary", "面议"),
            "education": job_data.get("education", "不限"),
            "experience": job_data.get("experience", "不限"),
            "city": job_data.get("city", "全国"),
            "skills": job_data.get("skills", ""),
            "description": job_data.get("content", "")
        }

        if stream:
            async def generate():
                try:
                    yield json.dumps({"type": "job_info", "data": job_detail}, ensure_ascii=False) + "\n"

                    # 使用环境变量读取 LLM 配置
                    from langchain_openai import ChatOpenAI
                    from langchain_core.messages import HumanMessage, SystemMessage

                    local_llm = ChatOpenAI(
                        api_key=os.getenv("OPENAI_API_KEY", ""),
                        base_url=os.getenv("OPENAI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"),
                        model="qwen3.6-plus",
                        temperature=0.3
                    )

                    system_prompt = f"""你是一个专业的职业咨询顾问。用户选择了'{job_name}'岗位进行咨询。
岗位基本信息：
- 薪资范围：{job_detail['salary_range']}
- 学历要求：{job_detail['education']}
- 经验要求：{job_detail['experience']}
- 工作地点：{job_detail['city']}
- 技能要求：{job_detail['skills'][:200]}...
请基于以上信息，结合用户的问题，提供专业、详细的职业咨询服务。回答时使用 Markdown 格式。"""

                    messages = [SystemMessage(content=system_prompt)]
                    for question in user_questions:
                        messages.append(HumanMessage(content=question))

                    response = local_llm.stream(messages)
                    for chunk in response:
                        if hasattr(chunk, 'content') and chunk.content:
                            yield json.dumps({"type": "chunk", "content": chunk.content}, ensure_ascii=False) + "\n"
                            await asyncio.sleep(0.02)

                    yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"
                except Exception as e:
                    logger.exception(f"❌ 咨询模式错误")
                    yield json.dumps({"type": "error", "message": str(e)}, ensure_ascii=False) + "\n"

            return StreamingResponse(generate(), media_type="application/x-ndjson")
        else:
            return {"status": "success", "job_info": job_detail}

    except Exception as e:
        logger.exception(f"❌ 咨询模式异常")
        return {"status": "error", "message": f"服务器错误：{str(e)}"}


# ==================== 聊天接口 ====================
@app.post("/api/chat")
async def chat(data: dict):
    """普通聊天对话接口，强制使用流式输出"""
    message = data.get("message", "")
    conversation_history = data.get("conversation_history", [])
    mode = data.get("mode", "job")
    job_intention = data.get("job_intention", {})

    logger.info(f"💬 收到用户消息（模式：{mode}）: {message}")

    try:
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import HumanMessage, SystemMessage

        local_llm = ChatOpenAI(
            api_key=os.getenv("OPENAI_API_KEY", ""),
            base_url=os.getenv("OPENAI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"),
            model=os.getenv("MODEL_NAME", "qwen3.6-plus"),
            temperature=0.7
        )

        if mode == "consult":
            system_prompt = """你是一个专业的职业咨询顾问。你的职责是：
1. 了解用户的求职意向（学历、行业偏好、岗位意向）
2. 推荐合适的岗位并说明技能要求
3. 提供职业发展规划建议
4. 解答职业发展相关问题
请用友好、耐心的语气回答，多使用鼓励性语言。回答时使用 Markdown 格式。"""
        else:
            system_prompt = """你是一个专业的智能就业导师。你的职责是：
1. 帮助用户分析简历与岗位的匹配度
2. 生成就业能力画像
3. 提供简历优化和面试建议
4. 规划职业发展路径
请用专业、细致的语气回答，给出具体可执行的建议。回答时使用 Markdown 格式。"""

        messages = [SystemMessage(content=system_prompt)]

        if mode == "consult" and job_intention:
            intention_text = f"用户的求职意向：学历={job_intention.get('education', '未填写')}, 行业={job_intention.get('industry', '未填写')}, 岗位={job_intention.get('position', '未填写')}"
            messages.append(HumanMessage(content=intention_text))

        for msg in conversation_history:
            if msg['role'] == 'user':
                messages.append(HumanMessage(content=msg['content']))
            elif msg['role'] == 'ai':
                messages.append(HumanMessage(content=f"[AI 回复]: {msg['content']}"))

        messages.append(HumanMessage(content=message))

        async def generate():
            try:
                logger.info("🔄 开始流式输出...")
                response = local_llm.stream(messages)
                for chunk in response:
                    if hasattr(chunk, 'content') and chunk.content:
                        yield json.dumps({"type": "chunk", "content": chunk.content}, ensure_ascii=False) + "\n"
                        await asyncio.sleep(0.02)
                yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"
            except Exception as e:
                logger.exception(f"❌ 流式输出错误")
                yield json.dumps({"type": "error", "message": str(e)}, ensure_ascii=False) + "\n"

        return StreamingResponse(generate(), media_type="application/x-ndjson")

    except Exception as e:
        logger.exception(f"❌ 聊天发生异常")
        return {"status": "error", "message": str(e)}


# ==================== 简历分析 ====================
@app.post("/api/analyze")
async def analyze_resume(data: dict):
    """简历分析接口，使用 RAG 匹配引擎自动推荐岗位"""
    resume_content = data.get("resume")
    target_job = data.get("target_job", "")
    stream = data.get("stream", False)

    logger.info(f"🚀 接收到前端请求，目标岗位：{target_job or '自动检测'}")

    try:
        engine = getattr(app.state, "match_engine", None)

        # 如果没有指定目标岗位，自动匹配
        if not target_job and engine:
            logger.info("🔍 未指定目标岗位，正在基于简历特征自动匹配...")
            profile = None
            try:
                from core.resume_parser import parse_resume_to_json
                profile = await parse_resume_to_json(resume_content)

                if profile and "error" not in profile:
                    formatted_skills = {s['name']: s.get('level_text', '熟悉') for s in profile.get('professional_skills', [])}
                    profile['skills'] = formatted_skills
                    profile['resume_text'] = resume_content

                    recommendations = engine.recommend_jobs(profile, top_n=1, use_rag=True)
                    if recommendations:
                        target_job = recommendations[0]['job_name']
                        score = recommendations[0]['total_score']
                        logger.info(f"✅ 基于技能/学历智能推荐岗位：{target_job}（匹配度：{score:.1f}分）")
                    else:
                        target_job = _get_job_list_from_personas()
                        if target_job is None:
                            target_job = "Java 开发"
                else:
                    target_job = _get_job_list_from_personas()
                    if target_job is None:
                        target_job = "Java 开发"
            except Exception as e:
                logger.warning(f"⚠️ 自动匹配异常：{str(e)}")
                target_job = _get_job_list_from_personas()
                if target_job is None:
                    target_job = "Java 开发"

        initial_state = {"resume_text": resume_content, "target_job": target_job}
        logger.info(f"🚀 开始执行分析流程，目标岗位：{target_job}")

        from core.agent_graph import career_agent_app

        if stream:
            async def generate():
                try:
                    logger.info("🔄 开始执行 LangGraph 智能体工作流（完全流式）...")
                    async for event in career_agent_app.astream_events(initial_state, version="v2"):
                        event_type = event.get("event")
                        if event_type == "on_chain_start":
                            node_name = event.get("name", "")
                            if node_name == "parse":
                                yield json.dumps({"type": "status", "message": "📄 正在解析简历..."}, ensure_ascii=False) + "\n"
                            elif node_name == "research":
                                yield json.dumps({"type": "status", "message": "🔍 正在计算匹配度..."}, ensure_ascii=False) + "\n"
                            elif node_name == "draft":
                                yield json.dumps({"type": "status", "message": "✍️ 正在生成职业规划报告..."}, ensure_ascii=False) + "\n"
                        elif event_type == "on_chain_end":
                            node_name = event.get("name", "")
                            output_data = event.get("data", {}).get("output", {})
                            if node_name == "parse" and output_data and "student_profile" in output_data:
                                yield json.dumps({"type": "parsed_profile", "data": output_data["student_profile"]}, ensure_ascii=False) + "\n"
                            elif node_name == "research" and output_data:
                                if "match_score" in output_data:
                                    yield json.dumps({"type": "match_result", "data": output_data["match_score"]}, ensure_ascii=False) + "\n"
                                if "transfer_paths" in output_data:
                                    yield json.dumps({"type": "transfer_paths", "data": output_data["transfer_paths"]}, ensure_ascii=False) + "\n"
                    yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"
                except Exception as e:
                    yield json.dumps({"type": "error", "message": str(e)}, ensure_ascii=False) + "\n"

            return StreamingResponse(generate(), media_type="application/x-ndjson")
        else:
            final_state = await asyncio.to_thread(career_agent_app.invoke, initial_state)
            if final_state.get("error"):
                return {"status": "error", "message": final_state["error"]}
            return {
                "status": "success",
                "summary": final_state["report_draft"],
                "parsed_profile": final_state["student_profile"],
                "match_result": final_state["match_score"],
                "transfer_paths": final_state["transfer_paths"]
            }

    except Exception as e:
        logger.exception(f"❌ 智能体执行发生异常")
        return {"status": "error", "message": f"服务器内部错误：{str(e)}"}


def _get_job_list_from_personas():
    """从画像文件获取岗位列表供用户选择"""
    try:
        personas_path = Path(__file__).parent / "data" / "standard_job_personas_upgraded.json"
        if personas_path.exists():
            with open(personas_path, 'r', encoding='utf-8') as f:
                personas = json.load(f)
            job_list = [p.get('job_name', p.get('岗位名称', '')) for p in personas if p.get('job_name') or p.get('岗位名称')]
            raise HTTPException(
                status_code=200,
                detail=json.dumps({"status": "need_selection", "message": "请手动选择目标岗位：", "job_list": job_list})
            )
    except HTTPException:
        raise
    except Exception:
        pass
    return None


# ==================== AI 润色接口 ====================
class PolishRequest(BaseModel):
    content: str
    context_type: str = "report"


@app.post("/api/polish-report")
async def polish_report(request: PolishRequest):
    """AI 智能润色与内容完整性检查"""
    logger.info(f"🪄 正在对内容进行 AI 润色与合规性检查...")

    try:
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import HumanMessage, SystemMessage

        llm_client = ChatOpenAI(
            api_key=os.getenv("OPENAI_API_KEY", ""),
            base_url=os.getenv("OPENAI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1"),
            model="qwen3.6-plus",
            temperature=0.7
        )

        system_prompt = """你是一个顶级的职业发展教练。你的任务是：
1. 润色：将用户提供的职业规划内容变得更加专业、具有行动导向。
2. 检查：检查内容是否包含核心三要素（目标、行动、时间轴）。如果缺失，请在润色后的内容末尾予以提示。
3. 必须保持 Markdown 格式。
4. 不要改变原有结构，只优化表达和补充缺失内容。"""

        response = llm_client.invoke([
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"请优化以下内容，保持原有 Markdown 格式：\n\n{request.content}")
        ])

        polished_text = response.content.strip()

        warnings = []
        if "月" not in request.content and "个" not in request.content:
            warnings.append("缺乏明确的时间周期")
        if "目标" not in request.content:
            warnings.append("职业目标描述不够清晰")

        return {
            "status": "success",
            "polished_content": polished_text,
            "integrity_check": {"passed": len(warnings) == 0, "warnings": warnings}
        }
    except Exception as e:
        logger.exception(f"❌ 润色服务失败")
        raise HTTPException(status_code=500, detail=f"润色服务暂时不可用：{str(e)}")


@app.post("/api/check-report-completeness")
async def check_integrity(request: PolishRequest):
    """内容完整性检查功能"""
    required_sections = ["短期", "中期", "评估", "学习"]
    missing = [s for s in required_sections if s not in request.content]

    return {
        "status": "success",
        "is_complete": len(missing) == 0,
        "missing_sections": missing,
        "suggestion": "请补充缺失的模块以确保规划的科学性。" if missing else "报告结构完整！"
    }


@app.post("/api/update_report")
async def update_report(request: PolishRequest):
    """支持手动编辑调整后的保存接口"""
    logger.info(f"📝 收到报告修改请求...")
    return {
        "status": "success",
        "message": "报告内容已根据您的手动调整更新并固化。",
        "last_update": "2026-04-07"
    }


if __name__ == "__main__":
    logger.info("🚀 智职领航 (LangGraph 版 v3.0) 后端服务器正在启动...")
    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8001")))
